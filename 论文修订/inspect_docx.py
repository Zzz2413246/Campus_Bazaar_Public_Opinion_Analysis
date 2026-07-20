from pathlib import Path
import sys
from docx import Document


source = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).with_name("原稿.docx")
doc = Document(source)

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\t", "\\t").replace("\n", "\\n")
    if text.strip():
        print(f"P{i:03d}\t[{paragraph.style.name}]\t{text}")

for table_index, table in enumerate(doc.tables):
    print(f"\nTABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        cells = [cell.text.replace("\t", "\\t").replace("\n", " / ") for cell in row.cells]
        print(f"R{row_index:02d}\t" + "\t".join(cells))
