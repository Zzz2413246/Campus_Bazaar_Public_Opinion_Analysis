from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "原稿.docx"
OUTPUT = BASE / "血常规参数联合CRP鉴别上呼吸道感染_学术规范修订版.docx"


paragraph_texts = {
    0: "血常规参数联合C反应蛋白鉴别细菌性与病毒性上呼吸道感染的回顾性分析",
    1: (
        "【摘要】目的  探讨血常规参数联合C反应蛋白（CRP）对细菌性与病毒性上呼吸道感染的辅助鉴别价值。"
        "方法  回顾性纳入2024年5月—2025年8月收治的300例上呼吸道感染患者，以病历记录中的最终病原学结论为参照，"
        "分为细菌组168例和病毒组132例。比较两组白细胞计数（WBC）、中性粒细胞百分比（NEUT%）、淋巴细胞百分比"
        "（LYMPH%）及CRP水平，并依据既有联合判定结果复核诊断效能。结果  细菌组WBC、NEUT%和CRP高于病毒组，"
        "LYMPH%低于病毒组，差异均有统计学意义（均P＜0.001）。既有联合判定的准确率、灵敏度和特异度分别为"
        "94.00%（282/300，95%置信区间：90.72%～96.17%）、95.24%（160/168，95%置信区间：90.89%～97.57%）"
        "和92.42%（122/132，95%置信区间：86.62%～95.83%）。结论  在本组病例及既定判定规则下，血常规参数联合"
        "CRP显示出较好的辅助鉴别效能。鉴于联合阈值、病原学平台和外部验证资料缺失，该结果尚不能作为单独决定抗菌"
        "治疗的依据，仍需结合临床表现和规范病原学检查解释。"
    ),
    2: "【关键词】上呼吸道感染；细菌感染；病毒感染；血常规；C反应蛋白；鉴别诊断",
    3: (
        "细菌性与病毒性上呼吸道感染常有相似的发热、咽痛和咳嗽表现，仅凭症状作出病原学判断容易产生偏差。"
        "病原培养、抗原检测或核酸检测能够提供较直接的病原学证据，但在部分门急诊场景中，检测可及性、周转时间及"
        "病原谱覆盖范围仍会限制其即时应用。WBC及白细胞分类计数获取便捷，CRP则反映炎症反应强度；两者均属于宿主"
        "反应指标，而非病原体特异性指标[1-2]。因此，临床真正需要回答的并不是某一指标是否“升高”，而是多项指标"
        "在明确参照标准下能否提供可重复的增量信息。基于既有病例汇总资料，本研究比较细菌性与病毒性上呼吸道感染"
        "患者的WBC、NEUT%、LYMPH%及CRP，并复核原联合判定结果的诊断效能，以评价其适用范围和局限。"
    ),
    4: "1  资料与方法",
    5: "1.1  研究对象",
    6: (
        "回顾性收集2024年5月—2025年8月收治的300例上呼吸道感染患者资料。患者年龄3～68岁，平均"
        "（35.12±4.26）岁；男158例，女142例；病程1～7 d，平均（3.25±1.05）d。依据病历记录中的最终病原学结论，"
        "细菌组168例，病毒组132例。纳入标准：符合上呼吸道感染的临床诊断要求；有发热、咽痛或咳嗽等相关表现；"
        "发病时间≤7 d；检测前未自行使用抗菌药物或抗病毒药物；临床及检验资料完整。排除标准：合并下呼吸道或其他"
        "部位感染；合并自身免疫性疾病、血液系统疾病或恶性肿瘤；严重肝肾功能不全；近期手术、创伤或明显应激；"
        "妊娠期或哺乳期。现有材料仅记载研究遵循医学伦理原则，未载明伦理审查机构、审批编号及知情同意豁免情况，"
        "正式投稿前须依据原始伦理文件补录。"
    ),
    7: "1.2  检测方法",
    8: (
        "采集外周静脉血2 mL，按检测项目要求分样。WBC、NEUT%和LYMPH%使用全自动血细胞分析仪"
        "（深圳市迈瑞生物医疗电子股份有限公司，BC-5385CRP）测定；CRP按现有资料记载采用干式免疫荧光法测定。"
        "原始材料列示的参考范围为WBC 4.0～10.0×10⁹/L、NEUT% 50%～70%、LYMPH% 20%～40%、CRP 0～10 mg/L。"
        "由于研究对象包含未成年人，结果解释原则上应采用本实验室经验证且与年龄、样本类型及检测系统相匹配的参考"
        "区间[2]。现有资料未提供CRP试剂厂家、批号、分析性能及室间质量评价结果，相关信息需在投稿前从检验系统"
        "原始记录中核实。"
    ),
    9: "1.3  参照标准与联合判定",
    10: (
        "以病历中的最终病原学检测结论作为参照标准：细菌性感染定义为阳性类别，病毒性感染定义为阴性类别。"
        "将既有联合判定结果与参照标准逐例对照，获得真阳性、假阳性、真阴性和假阴性例数。现有资料未记录病原学"
        "检测平台、采样部位、试剂信息以及联合判定的变量权重和阈值，故本研究仅对已形成的分组和效能结果进行复核，"
        "不补设缺乏原始依据的判定规则。"
    ),
    11: "1.4  观察指标",
    12: (
        "观察两组WBC、NEUT%、LYMPH%和CRP水平。联合判定效能以准确率、灵敏度、特异度、阳性预测值和阴性预测值"
        "表示。准确率=（真阳性+真阴性）/总例数×100%；灵敏度=真阳性/细菌性感染例数×100%；特异度=真阴性/"
        "病毒性感染例数×100%；阳性预测值=真阳性/联合判定为细菌性感染例数×100%；阴性预测值=真阴性/"
        "联合判定为病毒性感染例数×100%。"
    ),
    13: "1.5  统计学方法",
    14: (
        "采用SPSS 22.0进行统计分析。计量资料以均值±标准差表示；依据现有汇总数据，组间比较采用合并方差的两独立"
        "样本t检验。计数资料以例数和百分比表示，诊断效能的95%置信区间采用威尔逊法计算。检验均为双侧，P＜0.05"
        "为差异有统计学意义。因病例级数据不可得，本研究无法重新检验正态性、方差齐性、异常值及缺失数据处理过程。"
    ),
    15: "2  结果",
    16: "2.1  两组血常规参数及CRP水平",
    17: (
        "细菌组WBC、NEUT%和CRP均高于病毒组，LYMPH%低于病毒组。按汇总数据复核，四项指标的组间差异均达到"
        "统计学显著水平（均P＜0.001），见表1。"
    ),
    18: "表1  两组血常规参数及CRP水平比较（均值±标准差）",
    19: "2.2  联合判定与病原学结果的对应关系",
    20: (
        "既有联合判定结果中，170例被判为细菌性感染，其中160例与病原学结论一致，10例为假阳性；130例被判为"
        "病毒性感染，其中122例与病原学结论一致，8例为假阴性。具体构成见表2。"
    ),
    21: "表2  联合判定结果与病原学结论的四格表（例）",
    22: "2.3  联合判定的诊断效能",
    23: (
        "联合判定的准确率为94.00%，灵敏度为95.24%，特异度为92.42%；阳性预测值和阴性预测值分别为94.12%和"
        "93.85%。各指标的分子、分母及95%置信区间见表3。"
    ),
    24: "表3  血常规参数联合CRP的诊断效能",
    25: "3  讨论",
    26: (
        "本组资料中，细菌组的WBC、NEUT%和CRP较高，而LYMPH%较低。这一方向与感染后的宿主反应特征相符："
        "细菌感染可伴中性粒细胞反应增强，病毒感染时外周血细胞构成则受病毒种类、病程阶段和个体免疫状态共同影响。"
        "CRP由肝细胞在炎症介质刺激下合成，细菌感染时往往升高更明显，但创伤、自身免疫性炎症及其他非感染因素同样"
        "可使其升高[1-2]。上述差异说明这些指标具有分层价值，却不能据此推断其具有病原体特异性。"
    ),
    27: (
        "两组CRP均值分别为（35.62±8.56）mg/L和（6.35±2.12）mg/L，组间差异较大；WBC和白细胞分类也呈一致"
        "方向。既往研究同样观察到CRP、WBC或其他炎症指标联合分析能够增加呼吸道感染评估的信息量[3-5]。然而，"
        "本研究只有汇总均值和标准差，无法展示个体数值的重叠程度，也不能据此推导最佳截断值。换言之，组间均值"
        "差异显著并不等同于单例患者能够被无误分类，临床解释仍需回到具体病程和检测背景。"
    ),
    28: (
        "既有联合判定在本样本中的灵敏度为95.24%，特异度为92.42%，表面上具有较好的区分能力，但其适用边界必须"
        "结合研究设计理解。首先，联合规则的阈值及建模过程未被记录，其他机构无法据此复现；其次，效能指标来自同一"
        "批病例，未经过独立验证，可能高估真实应用表现；再次，阳性预测值和阴性预测值会随细菌性感染比例改变，"
        "本研究结果不能直接移用于病原谱和就诊人群不同的场景。相关随机试验对CRP辅助抗菌药物决策的结果亦不完全"
        "一致，提示检测的价值取决于应用流程、阈值设置和临床情境[6-7]。"
    ),
    29: (
        "对首次结果位于重叠区间、症状持续或病情进展的患者，更合理的做法是结合病程复查血常规和CRP，并按临床需要"
        "完善抗原、核酸或培养等病原学检查。我国流行性感冒和人偏肺病毒感染诊疗方案均强调结合临床表现、流行病学"
        "信息及病原学证据进行诊断[8-9]。因此，血常规与CRP更适合作为快速风险分层的一部分，而不宜被解释为启动、"
        "维持或排除抗菌治疗的单一门槛。"
    ),
    30: (
        "本研究存在若干限制。其一，缺少病例级数据，无法核查分布假设、异常值、缺失值及潜在混杂因素；其二，病原学"
        "平台、采样部位和联合判定阈值未记录，研究可重复性受限；其三，资料来自单一来源，且未设置内部验证或外部"
        "验证；其四，研究对象年龄跨度为3～68岁，未按儿童和成人分层，不同年龄参考区间及免疫反应差异可能影响结果；"
        "其五，伦理审批信息及部分试剂质量控制资料缺失。正式投稿前，应依据原始病历、检验信息系统和伦理文件逐项"
        "补齐，而不应以推测性内容替代。"
    ),
    31: (
        "综上，本组细菌性与病毒性上呼吸道感染患者的WBC、NEUT%、LYMPH%及CRP水平存在统计学差异，既有联合判定"
        "在本样本中表现出较高的灵敏度和特异度。由于关键阈值和验证信息尚不完整，目前证据仅支持其用于辅助判断和"
        "风险分层。后续研究应预先明确病原学参照标准与联合算法，并在年龄分层和独立样本中验证后，再评价其临床推广价值。"
    ),
    32: "参考文献：",
}

references = [
    "[1] 葛均波，王辰，王建安. 内科学[M]. 10版. 北京：人民卫生出版社，2024.",
    "[2] 中华人民共和国国家卫生健康委员会. WS/T 404.9—2018 临床常用生化检验项目参考区间 第9部分：血清C-反应蛋白、前白蛋白、转铁蛋白、β2-微球蛋白[S]. 2018.",
    "[3] 左晶，陈海鹰. SAA、CRP及血常规检测在急性病毒性呼吸道感染中的临床诊断价值[J]. 生命科学仪器，2025，23（6）：25-27.",
    "[4] 练德峰，王勇. 血清降钙素原、C反应蛋白与白细胞总数联合检测在感染疾病中的诊断应用[J]. 湖北中医杂志，2015，37（11）：18-19.",
    "[5] 张晓敏，陈清勇. 降钙素原、C反应蛋白在社区获得性肺炎细菌感染中的诊断价值[J]. 全科医学临床与教育，2014，12（2）：138-141.",
    "[6] Do N T T, Ta N T D, Tran N T H, et al. Implementation of point-of-care testing of C-reactive protein concentrations to improve antibiotic targeting in respiratory illness in Vietnamese primary care: a pragmatic cluster-randomised controlled trial[J]. Lancet Infectious Diseases, 2023, 23(9): 1085-1094. DOI:10.1016/S1473-3099(23)00125-1.",
    "[7] Jung C, Levy C, Béchet S, et al. Impact of C-reactive protein point-of-care testing on antibiotic prescriptions for children and adults with suspected respiratory tract infections in primary care: a patient-level randomized controlled trial[J]. Clinical Microbiology and Infection, 2024, 30(12): 1553-1558. DOI:10.1016/j.cmi.2024.07.014.",
    "[8] 中华人民共和国国家卫生健康委员会，国家中医药管理局. 流行性感冒诊疗方案（2018年版修订版）[S]. 2018.",
    "[9] 中华人民共和国国家卫生健康委员会，国家中医药局. 人偏肺病毒感染诊疗方案（2023年版）[S]. 2023.",
]


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=10.5, bold=None):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9)


doc = Document(SOURCE)
original_paragraphs = list(doc.paragraphs)

for index, text in paragraph_texts.items():
    original_paragraphs[index].text = text

for paragraph in original_paragraphs[33:]:
    paragraph._element.getparent().remove(paragraph._element)

for reference in references:
    doc.add_paragraph(reference)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.4)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

normal_style = doc.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(10.5)
normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.space_before = Pt(0)
normal_style.paragraph_format.space_after = Pt(0)

for paragraph in doc.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.keep_together = False
    for run in paragraph.runs:
        set_run_font(run)

title = doc.paragraphs[0]
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.first_line_indent = Pt(0)
title.paragraph_format.space_after = Pt(10)
title.paragraph_format.line_spacing = 1.2
for run in title.runs:
    set_run_font(run, east_asia="黑体", size=16, bold=True)

for index in (1, 2):
    paragraph = doc.paragraphs[index]
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(5 if index == 1 else 8)

for index in (4, 15, 25, 32):
    paragraph = doc.paragraphs[index]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_font(run, east_asia="黑体", size=12, bold=True)

for index in (5, 7, 9, 11, 13, 16, 19, 22):
    paragraph = doc.paragraphs[index]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_font(run, east_asia="黑体", size=10.5, bold=True)

for index in (18, 21, 24):
    paragraph = doc.paragraphs[index]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_font(run, size=10, bold=False)

reference_start = len(doc.paragraphs) - len(references)
for paragraph in doc.paragraphs[reference_start:]:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(21)
    paragraph.paragraph_format.first_line_indent = Pt(-21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        set_run_font(run, size=9.5)

for table in doc.tables:
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    set_run_font(run, size=8.5, bold=(row_index == 0))

footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.clear()
add_page_number(footer_paragraph)

core = doc.core_properties
core.title = paragraph_texts[0]
core.subject = "血常规参数联合CRP辅助鉴别细菌性与病毒性上呼吸道感染"
core.keywords = "上呼吸道感染; 血常规; C反应蛋白; 细菌感染; 病毒感染"
core.comments = "在不补造原始信息的前提下完成学术规范修订。"

doc.save(OUTPUT)
print(OUTPUT)
