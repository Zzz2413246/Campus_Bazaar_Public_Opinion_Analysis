from docx import Document
from pathlib import Path

doc_path = Path(__file__).with_name("参考报告.docx")
out_path = Path(__file__).with_name("参考报告全文.txt")
doc = Document(doc_path)

lines = []
for i, paragraph in enumerate(doc.paragraphs):
    lines.append(f"P{i}: {paragraph.text}")

lines.append(f"\nTABLES: {len(doc.tables)}")
for table_index, table in enumerate(doc.tables):
    lines.append(f"\nTABLE {table_index}")
    for row in table.rows:
        lines.append(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))

out_path.write_text("\n".join(lines), encoding="utf-8")
print(out_path)
