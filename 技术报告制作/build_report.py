from __future__ import annotations

import math
import os
import zipfile
from xml.etree import ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).parent
REFERENCE = ROOT / "参考报告.docx"
OUTPUT = ROOT / "血常规指标及CRP鉴别细菌性与病毒性上呼吸道感染_规范技术报告.docx"


def set_run_font(run, size=12, bold=None, latin="Times New Roman"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_base(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=True,
    keep_with_next=False,
):
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = keep_with_next
    pf.widow_control = True
    if first_line:
        pf.first_line_indent = Inches(0.333)
    else:
        pf.first_line_indent = None


def add_body(doc, text, *, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    set_paragraph_base(p, alignment=alignment, first_line=True)
    r = p.add_run(text)
    set_run_font(r, 12, bold)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_base(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=True,
        keep_with_next=True,
    )
    r = p.add_run(text)
    set_run_font(r, 12, True)
    return p


def add_labeled_paragraph(doc, pieces, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    set_paragraph_base(p, alignment=alignment, first_line=False)
    for text, bold in pieces:
        r = p.add_run(text)
        set_run_font(r, 12, bold)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_base(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        keep_with_next=True,
    )
    r = p.add_run(text)
    set_run_font(r, 12, False)
    return p


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, top=True, header=True, bottom=True):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        enabled = (edge == "top" and top) or (edge == "bottom" and bottom)
        node.set(qn("w:val"), "single" if enabled else "nil")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:color"), "000000")
    if header and table.rows:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        trPr.append(tblHeader)
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            bottom_node = OxmlElement("w:bottom")
            bottom_node.set(qn("w:val"), "single")
            bottom_node.set(qn("w:sz"), "8")
            bottom_node.set(qn("w:color"), "000000")
            tcBorders.append(bottom_node)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def fill_table(table, rows, widths, *, font_size=10.5):
    for row_index, values in enumerate(rows):
        cells = table.rows[row_index].cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, font_size, row_index == 0)
    set_table_geometry(table, widths)
    set_table_borders(table)


def wilson(x, n, z=1.959963984540054):
    p = x / n
    d = 1 + z * z / n
    c = (p + z * z / (2 * n)) / d
    h = z * math.sqrt(p * (1 - p) / n + z * z / (4 * n * n)) / d
    return p * 100, (c - h) * 100, (c + h) * 100


def clean_unused_template_parts(path):
    """Remove legacy OLE/data objects that no longer have body references."""
    temp_path = path.with_name(path.stem + ".cleaning.docx")
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    skip_prefixes = ("word/embeddings/", "word/media/", "customXml/")
    skip_exact = {"word/people.xml"}
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            name = item.filename
            if name in skip_exact or any(name.startswith(prefix) for prefix in skip_prefixes):
                continue
            data = source.read(name)
            if name == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                for rel in list(root):
                    rel_type = rel.attrib.get("Type", "")
                    if rel_type.endswith(("/oleObject", "/image", "/customXml", "/people")):
                        root.remove(rel)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "[Content_Types].xml":
                root = ET.fromstring(data)
                for node in list(root):
                    part_name = node.attrib.get("PartName", "")
                    extension = node.attrib.get("Extension", "")
                    if (
                        part_name.startswith(("/word/embeddings/", "/word/media/", "/customXml/"))
                        or part_name == "/word/people.xml"
                        or extension in {"bin", "wmf"}
                    ):
                        root.remove(node)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, data)
    os.replace(temp_path, path)


# Recreate the distilled template in a clean OOXML package. The reference file
# contains legacy embedded WPS/Word OLE objects that make Word automation hang
# after body replacement, so only its verified page and typography tokens are
# carried forward.
doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.268)
section.page_height = Inches(11.693)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)
section.header_distance = Inches(0.591)
section.footer_distance = Inches(0.689)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

doc.core_properties.title = "血常规指标及CRP在鉴别诊断细菌性与病毒性上呼吸道感染中的应用分析"
doc.core_properties.subject = "临床检验技术报告"
doc.core_properties.author = ""
doc.core_properties.keywords = "上呼吸道感染；细菌性；病毒性；血常规；C反应蛋白"

title = doc.add_paragraph()
set_paragraph_base(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
title.paragraph_format.space_after = Pt(12)
run = title.add_run("血常规指标及CRP在鉴别诊断细菌性与病毒性\n上呼吸道感染中的应用分析")
set_run_font(run, 16, True)

add_labeled_paragraph(
    doc,
    [
        ("【摘要】", True),
        ("目的  ", True),
        ("评价血常规相关指标联合C反应蛋白检测对细菌性与病毒性上呼吸道感染的辅助鉴别价值。", False),
        ("方法  ", True),
        ("回顾性整理2024年5月—2025年8月收治的300例上呼吸道感染患者资料，以最终病原学检测结论为参照，分为细菌组168例和病毒组132例。比较两组白细胞计数、中性粒细胞百分比、淋巴细胞百分比及C反应蛋白水平，并评价联合检测的诊断效能。", False),
        ("结果  ", True),
        ("细菌组白细胞计数、中性粒细胞百分比和C反应蛋白水平均高于病毒组，淋巴细胞百分比低于病毒组，差异均有统计学意义（P＜0.001）。联合检测准确率为94.00%（282/300，95%置信区间：90.72%～96.17%），灵敏度为95.24%（160/168，95%置信区间：90.89%～97.57%），特异度为92.42%（122/132，95%置信区间：86.62%～95.83%）。", False),
        ("结论  ", True),
        ("血常规指标联合C反应蛋白检测具有较好的辅助鉴别价值，可为临床分层评估和合理用药提供实验室依据，但不能替代规范的病原学检测及综合临床判断。", False),
    ],
)

add_labeled_paragraph(
    doc,
    [
        ("【关键词】", True),
        ("上呼吸道感染；细菌感染；病毒感染；血常规；C反应蛋白；鉴别诊断", False),
    ],
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
)

add_body(
    doc,
    "上呼吸道感染是临床常见的急性呼吸道疾病，病毒和细菌均可致病。两类感染在发热、咽痛、咳嗽等症状上存在重叠，但抗感染治疗策略不同，因此早期判断病原学倾向具有现实意义[1]。病原培养、抗原或核酸检测可提供病原学证据，但在基层或急诊场景中可能受检测周期、可及性和病原谱覆盖范围限制。血常规检查周转快、成本低，WBC、NEUT%和LYMPH%能够反映外周血免疫细胞构成；CRP为急性时相反应蛋白，其结果应结合检测系统参考区间解释[2]。国内研究提示，CRP与血常规参数联合分析可为细菌性和病毒性呼吸道感染的初步鉴别提供辅助信息[3-7]。本报告依据既有病例资料，对血常规联合CRP的应用价值进行规范分析。",
)

add_heading(doc, "1  资料与方法")
add_heading(doc, "1.1  一般资料")
add_body(
    doc,
    "纳入2024年5月—2025年8月收治的300例上呼吸道感染患者。年龄3～68岁，平均（35.12±4.26）岁；男158例，女142例；病程1～7 d，平均（3.25±1.05）d。根据病原学检测结论分为细菌组168例和病毒组132例。纳入标准：符合上呼吸道感染临床诊断要求；具有发热、咽痛、咳嗽等相应表现；发病时间≤7 d；检测前未自行使用抗菌药物或抗病毒药物；临床与检验资料完整。排除标准：合并下呼吸道感染或其他部位感染；合并自身免疫性疾病、血液系统疾病或恶性肿瘤；严重肝肾功能不全；近期有手术、创伤或明显应激；妊娠期或哺乳期。原始材料说明研究符合医学伦理原则，但未提供伦理审批编号，正式用于投稿或成果鉴定前应补充核验。",
)

add_heading(doc, "1.2  检测方法")
add_body(
    doc,
    "采集外周静脉血2 mL并按检测要求分样。使用全自动血细胞分析仪（深圳市迈瑞生物医疗电子股份有限公司，型号BC-5385CRP，原材料登记信息：粤械注准20202220769）检测WBC、NEUT%和LYMPH%。CRP采用干式免疫荧光法检测。原材料所列参考范围为：WBC 4.0～10.0×10⁹/L、NEUT% 50%～70%、LYMPH% 20%～40%、CRP 0～10 mg/L。检测由检验专业人员依仪器操作规程和试剂说明书完成，并实施室内质量控制。不同检测系统的参考区间可能存在差异，结果解释应以本实验室经验证的参考区间为准[2]。",
)

add_heading(doc, "1.3  分组与联合判定")
add_body(
    doc,
    "以医疗记录中的最终病原学检测结论作为参照，将细菌性感染定义为阳性类别、病毒性感染定义为阴性类别。联合检测结果与参照标准逐例比较，形成真阳性、假阴性、真阴性和假阳性四格表。原始材料未记录具体病原学平台、试剂批号及联合判定阈值，故本报告仅复核其已形成的分组和效能数据，不对判定规则作未经资料支持的补充。",
)

add_heading(doc, "1.4  观察指标")
add_body(
    doc,
    "比较两组WBC、NEUT%、LYMPH%和CRP。诊断效能指标包括准确率、灵敏度、特异度、阳性预测值和阴性预测值。计算式分别为：准确率=（真阳性例数+真阴性例数）/总例数×100%；灵敏度=真阳性例数/细菌性感染总例数×100%；特异度=真阴性例数/病毒性感染总例数×100%；阳性预测值=真阳性例数/联合检测判断为细菌性感染的总例数×100%；阴性预测值=真阴性例数/联合检测判断为病毒性感染的总例数×100%。",
)

add_heading(doc, "1.5  统计学方法")
add_body(
    doc,
    "采用SPSS 22.0进行统计分析。计量资料以均值±标准差表示，两独立样本比较采用独立样本t检验；计数资料以例数和百分比表示。双侧P＜0.05为差异有统计学意义。诊断效能的95%置信区间采用威尔逊法计算。",
)

add_heading(doc, "2  结果")
add_heading(doc, "2.1  两组血常规及CRP指标比较")
add_body(
    doc,
    "细菌组WBC、NEUT%和CRP均高于病毒组，LYMPH%低于病毒组，四项指标组间差异均有统计学意义（P＜0.001），见表1。",
)
add_caption(doc, "表1  两组血常规及CRP指标比较（均值±标准差）")
table1_rows = [
    ["感染类型", "n", "WBC（×10⁹/L）", "NEUT%（%）", "LYMPH%（%）", "CRP（mg/L）"],
    ["细菌组", "168", "13.56±2.12", "78.25±5.36", "18.36±4.12", "35.62±8.56"],
    ["病毒组", "132", "6.85±1.53", "55.36±4.85", "36.52±5.23", "6.35±2.12"],
    ["t值", "—", "30.629", "38.273", "33.644", "38.360"],
    ["P值", "—", "＜0.001", "＜0.001", "＜0.001", "＜0.001"],
]
table1 = doc.add_table(rows=len(table1_rows), cols=6)
fill_table(table1, table1_rows, [1100, 520, 1750, 1500, 1550, 1880], font_size=9.5)

add_heading(doc, "2.2  联合检测结果")
add_body(
    doc,
    "血常规联合CRP判断为细菌性感染170例，其中160例与病原学结论一致、10例为假阳性；判断为病毒性感染130例，其中122例与病原学结论一致、8例为假阴性，见表2。",
)
add_caption(doc, "表2  血常规联合CRP与病原学检测结果的四格表（例）")
table2_rows = [
    ["联合检测结果", "病原学：细菌性", "病原学：病毒性", "合计"],
    ["细菌性", "160（真阳性）", "10（假阳性）", "170"],
    ["病毒性", "8（假阴性）", "122（真阴性）", "130"],
    ["合计", "168", "132", "300"],
]
table2 = doc.add_table(rows=len(table2_rows), cols=4)
fill_table(table2, table2_rows, [2000, 2300, 2300, 1700], font_size=10.5)

heading_23 = add_heading(doc, "2.3  联合检测的诊断效能")
heading_23.paragraph_format.page_break_before = True
metrics = [
    ("准确率", 282, 300),
    ("灵敏度", 160, 168),
    ("特异度", 122, 132),
    ("阳性预测值", 160, 170),
    ("阴性预测值", 122, 130),
]
table3_rows = [["指标", "分子/分母", "估计值（%）", "95%置信区间（%）"]]
for label, x, n in metrics:
    value, low, high = wilson(x, n)
    table3_rows.append([label, f"{x}/{n}", f"{value:.2f}", f"{low:.2f}～{high:.2f}"])
add_body(
    doc,
    "联合检测准确率为94.00%，灵敏度为95.24%，特异度为92.42%；阳性预测值为94.12%，阴性预测值为93.85%，见表3。",
)
add_caption(doc, "表3  血常规联合CRP的诊断效能")
table3 = doc.add_table(rows=len(table3_rows), cols=4)
fill_table(table3, table3_rows, [2100, 1800, 1900, 2500], font_size=10.5)
for row in table3.rows:
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)
    for cell in row.cells:
        set_cell_margins(cell, top=25, start=100, bottom=25, end=100)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0

add_heading(doc, "3  讨论")
add_body(
    doc,
    "本组资料显示，细菌组WBC、NEUT%和CRP明显高于病毒组，而LYMPH%较低。细菌感染常伴中性粒细胞相关免疫反应增强，外周血白细胞及中性粒细胞比例可升高；病毒感染的血细胞变化受病毒种类、病程和宿主免疫状态影响，部分病例可表现为相对淋巴细胞增多。CRP由肝细胞在炎症介质刺激下合成，细菌感染时通常升高更明显，但并不具有病原体特异性[1-2]。",
)
add_body(
    doc,
    "本资料中细菌组CRP为（35.62±8.56）mg/L，病毒组为（6.35±2.12）mg/L，差异显著。该结果与国内有关血常规、CRP及其他炎症指标在呼吸道感染鉴别中的研究方向基本一致[3-7]。但不同研究的病例构成、检测系统和判定标准并不相同，指标分布存在一定重叠。因此，CRP不能以单一固定界值替代病原学诊断。",
)
add_body(
    doc,
    "对于首次检测结果处于重叠区间、症状持续或病情进展者，应结合病程复查血常规和CRP，并完善抗原、核酸或培养等病原学检查。我国流行性感冒及人偏肺病毒感染相关诊疗方案均强调结合临床表现与病原学证据进行诊断和鉴别[8-9]。因此，连续观察指标变化通常比孤立解读单次结果更符合临床实际。",
)
add_body(
    doc,
    "本资料的联合检测准确率、灵敏度和特异度分别为94.00%、95.24%和92.42%，说明在该样本及既定判定规则下表现较好。国内相关研究也表明，将CRP与WBC、白细胞分类或其他炎症指标联合分析，可在一定程度上提高感染类型判断的参考价值[3-7]。但不同研究使用的检测平台、目标人群和参照标准不同，其效能数值不能直接外推。临床应用时，应把血常规和CRP作为风险分层工具，与症状体征、病程、流行病学信息、影像学及必要的培养或核酸检测共同判断。",
)
add_body(
    doc,
    "本报告存在以下局限：第一，原始材料未提供病例级数据，无法核查异常值、正态性、方差齐性及缺失数据处理；第二，未记录具体病原学检测平台、联合判定阈值及盲法流程，可能影响可重复性；第三，为单一资料来源的回顾性分析，未进行内部或外部验证；第四，纳入年龄跨度较大，未作年龄分层，儿童与成人参考区间及免疫反应差异可能造成混杂；第五，伦理审批编号未载明。上述信息应在正式投稿、职称评审或成果鉴定前依据原始病历、检验系统和伦理文件补齐。",
)

add_body(
    doc,
    "综上，细菌性与病毒性上呼吸道感染患者的WBC、NEUT%、LYMPH%及CRP分布存在显著差异，血常规联合CRP在本组资料中表现出较好的辅助鉴别效能。该组合适合用于快速初筛和临床分层，但不应作为单独启动或排除抗菌治疗的唯一依据，仍需结合规范病原学检测和整体临床证据。",
)

add_heading(doc, "参考文献：")
references = [
    "[1] 葛均波，徐永健，王辰. 内科学（第9版）[M]. 北京：人民卫生出版社，2018.",
    "[2] 中华人民共和国国家卫生健康委员会. WS/T 404.9—2018 临床常用生化检验项目参考区间 第9部分：血清C反应蛋白、前白蛋白、转铁蛋白、β2-微球蛋白[S]. 2018.",
    "[3] 陈苗苗，阙国勇，陈丹丹，等. 白细胞计数、降钙素原及血清淀粉样蛋白A联合检测在孕妇上呼吸道病毒性感染诊断中的应用价值[J]. 中国妇幼保健，2024，39（24）：4823-4827.",
    "[4] 徐鑫鑫，蔡花，李海泉，等. 南通地区甲型流感患儿血常规及CRP检测指标的临床价值分析[J]. 标记免疫分析与临床，2025，32（11）：2281-2286.",
    "[5] 左晶，陈海鹰. SAA、CRP及血常规检测在急性病毒性呼吸道感染中的临床诊断价值[J]. 生命科学仪器，2025，23（6）：25-27.",
    "[6] 练德峰，王勇. 血清降钙素原、C反应蛋白与白细胞总数联合检测在感染疾病中的诊断应用[J]. 湖北中医杂志，2015，37（11）：18-19.",
    "[7] 张晓敏，陈清勇. 降钙素原、C反应蛋白在社区获得性肺炎细菌感染中的诊断价值[J]. 全科医学临床与教育，2014，12（2）：138-141.",
    "[8] 中华人民共和国国家卫生健康委员会. 流行性感冒诊疗方案（2018年版修订版）[S]. 2018.",
    "[9] 中华人民共和国国家卫生健康委员会，国家中医药局. 人偏肺病毒感染诊疗方案（2023年版）[S]. 2023.",
]
for ref in references:
    p = doc.add_paragraph()
    set_paragraph_base(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.25)
    r = p.add_run(ref)
    set_run_font(r, 10.5, False)

settings = doc.settings._element
for tag in ("w:updateFields", "w:trackRevisions"):
    node = settings.find(qn(tag))
    if node is not None:
        settings.remove(node)

doc.save(OUTPUT)
clean_unused_template_parts(OUTPUT)
print(OUTPUT)
