import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document

root = Path(__file__).parent
path = root / "参考报告.docx"
doc = Document(path)

def inches(value):
    return None if value is None else round(value.inches, 3)

def pt(value):
    return None if value is None else round(value.pt, 2)

paragraphs = []
for index, p in enumerate(doc.paragraphs):
    first = next((r for r in p.runs if r.text.strip()), p.runs[0] if p.runs else None)
    paragraphs.append({
        "index": index,
        "text": p.text[:120],
        "style": p.style.name,
        "alignment": str(p.alignment),
        "left_indent_in": inches(p.paragraph_format.left_indent),
        "first_line_indent_in": inches(p.paragraph_format.first_line_indent),
        "space_before_pt": pt(p.paragraph_format.space_before),
        "space_after_pt": pt(p.paragraph_format.space_after),
        "line_spacing": p.paragraph_format.line_spacing,
        "keep_with_next": p.paragraph_format.keep_with_next,
        "font": first.font.name if first else None,
        "font_size_pt": pt(first.font.size) if first else None,
        "bold": first.bold if first else None,
    })

parts = []
with zipfile.ZipFile(path) as zf:
    for item in sorted(zf.infolist(), key=lambda x: x.filename):
        data = zf.read(item.filename)
        parts.append({
            "name": item.filename,
            "size": len(data),
            "sha256": hashlib.sha256(data).hexdigest(),
        })

section = doc.sections[0]
evidence = {
    "reference_sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
    "reference_size": path.stat().st_size,
    "section_count": len(doc.sections),
    "page": {
        "width_in": inches(section.page_width),
        "height_in": inches(section.page_height),
        "left_margin_in": inches(section.left_margin),
        "right_margin_in": inches(section.right_margin),
        "top_margin_in": inches(section.top_margin),
        "bottom_margin_in": inches(section.bottom_margin),
        "header_distance_in": inches(section.header_distance),
        "footer_distance_in": inches(section.footer_distance),
    },
    "paragraphs": paragraphs,
    "table_count": len(doc.tables),
    "table_grids_in": [
        [round(cell.width.inches, 3) if cell.width else None for cell in table.rows[0].cells]
        for table in doc.tables
    ],
    "parts": parts,
}

(root / "template_evidence.json").write_text(
    json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(root / "template_evidence.json")
